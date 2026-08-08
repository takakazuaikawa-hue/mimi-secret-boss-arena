import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SOURCE = Path(sys.argv[1])
OUTPUT = Path(sys.argv[2])

INK = RGBColor(38, 54, 74)
ACCENT = RGBColor(162, 59, 94)
ACCENT_DARK = RGBColor(111, 37, 66)
MUTED = RGBColor(105, 112, 122)
PALE = "F8EDF2"
RULE = "D8C2CB"
FONT = "Yu Gothic"


def set_run_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_or_paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)


def set_paragraph_border(paragraph, color=RULE, size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instr, separate, text, end):
        run._r.append(element)
    set_run_font(run, size=8.5, color=MUTED)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (17, ACCENT_DARK, 18, 10),
        "Heading 2": (14, ACCENT, 14, 7),
        "Heading 3": (11.5, ACCENT_DARK, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_run = header.add_run(
        "ミミのときめき裏ボス闘技場  |  テキスト編集用マスター"
    )
    set_run_font(header_run, size=8.5, color=MUTED, bold=True)
    set_paragraph_border(header, color="E5D8DE", size="5")

    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def add_cover(doc, data):
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(76)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("GAME NARRATIVE EDITING MANUSCRIPT")
    set_run_font(run, size=10, color=ACCENT, bold=True)
    kicker.paragraph_format.space_after = Pt(16)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(data["projectTitle"])
    set_run_font(title_run, size=28, color=INK, bold=True)
    title.paragraph_format.space_after = Pt(8)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("全物語テキスト 編集用マスター")
    set_run_font(subtitle_run, size=15, color=ACCENT_DARK, bold=True)
    subtitle.paragraph_format.space_after = Pt(48)

    warning = doc.add_paragraph()
    warning.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_cell_or_paragraph_shading(warning, PALE)
    warning.paragraph_format.left_indent = Inches(0.18)
    warning.paragraph_format.right_indent = Inches(0.18)
    warning.paragraph_format.space_before = Pt(10)
    warning.paragraph_format.space_after = Pt(10)
    warning.paragraph_format.line_spacing = 1.25
    label = warning.add_run("編集時のお願い\n")
    set_run_font(label, size=11, color=ACCENT_DARK, bold=True)
    body = warning.add_run(
        "灰色の角括弧ID（例：[CHAR.gidonozeaas.meet.L001]）は変更・削除しないでください。"
        "IDの後ろにある本文、話者名、選択肢、結果文は自由に書き換えられます。"
        "編集後のこのWordファイルを渡していただければ、IDを基準にゲームへ反映します。"
    )
    set_run_font(body, size=10.5, color=INK)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note_run = note.add_run(
        "現在ゲームに入っている文章をそのまま収録しています。旧設定や矛盾もレビュー対象として残しています。"
    )
    set_run_font(note_run, size=9, color=MUTED, italic=True)

    doc.add_page_break()


def add_id_text_paragraph(
    doc,
    stable_id,
    text,
    speaker=None,
    kind=None,
    shaded=False,
    indent=0,
):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(indent)
    paragraph.paragraph_format.first_line_indent = Inches(0)
    paragraph.paragraph_format.keep_together = True
    if shaded:
        set_cell_or_paragraph_shading(paragraph, "FCF6F8")
        paragraph.paragraph_format.left_indent = Inches(indent + 0.12)
        paragraph.paragraph_format.right_indent = Inches(0.12)
        paragraph.paragraph_format.space_before = Pt(4)
        paragraph.paragraph_format.space_after = Pt(7)

    id_run = paragraph.add_run(f"[{stable_id}] ")
    set_run_font(id_run, size=7.5, color=MUTED, bold=True)

    if speaker:
        speaker_label = f"{speaker}："
    elif kind == "thought":
        speaker_label = "ミミ（心の声）："
    else:
        speaker_label = "地の文："
    speaker_run = paragraph.add_run(speaker_label)
    set_run_font(speaker_run, size=10, color=ACCENT_DARK, bold=True)

    text_run = paragraph.add_run(text or "")
    set_run_font(text_run, size=10.8, color=INK)
    return paragraph


def add_metadata(doc, stable_id, title, value):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.keep_together = True
    id_run = paragraph.add_run(f"[{stable_id}] ")
    set_run_font(id_run, size=7.5, color=MUTED, bold=True)
    label_run = paragraph.add_run(f"{title}：")
    set_run_font(label_run, size=9.5, color=ACCENT_DARK, bold=True)
    value_run = paragraph.add_run(str(value or ""))
    set_run_font(value_run, size=10.5, color=INK)


def add_scene(doc, scene, prefix, heading_level=2):
    doc.add_heading(scene.get("title") or scene.get("id") or prefix, level=heading_level)
    add_metadata(doc, f"{prefix}.META.ID", "場面ID", scene.get("id", ""))
    add_metadata(doc, f"{prefix}.META.LOCATION", "場所", scene.get("location", ""))

    for index, line in enumerate(scene.get("lines", []), start=1):
        line_id = f"{prefix}.L{index:03d}"
        add_id_text_paragraph(
            doc,
            line_id,
            line.get("text", ""),
            speaker=line.get("speaker"),
            kind=line.get("kind"),
        )
        if line.get("cue"):
            add_metadata(doc, f"{line_id}.CUE", "補助文", line["cue"])

    choices = scene.get("choices") or []
    if choices:
        doc.add_heading("選択肢", level=3)
    for index, choice in enumerate(choices, start=1):
        choice_id = f"{prefix}.C{index:02d}"
        add_metadata(doc, f"{choice_id}.LABEL", f"選択肢{index}", choice.get("label", ""))
        add_metadata(doc, f"{choice_id}.RESULT", "選択後の本文", choice.get("result", ""))
        if choice.get("intent"):
            add_metadata(doc, f"{choice_id}.INTENT", "意図", choice["intent"])
        if choice.get("promise"):
            add_metadata(doc, f"{choice_id}.PROMISE", "選ぶ前の補助文", choice["promise"])
        if choice.get("memory"):
            add_metadata(doc, f"{choice_id}.MEMORY", "後に残る記憶", choice["memory"])
        mechanics = []
        for key in (
            "trust",
            "ownership",
            "money",
            "sharedPoints",
            "fighterPoints",
            "condition",
            "recruitmentDecision",
            "liberationDecision",
            "tone",
        ):
            value = choice.get(key)
            if value not in (None, 0, ""):
                mechanics.append(f"{key}={value}")
        if mechanics:
            note = doc.add_paragraph()
            note.paragraph_format.left_indent = Inches(0.18)
            run = note.add_run("参考（ゲーム処理）：" + " / ".join(mechanics))
            set_run_font(run, size=8.5, color=MUTED, italic=True)


def add_prologue(doc, pages, prefix, title):
    doc.add_heading(title, level=1)
    for index, page in enumerate(pages, start=1):
        stable_id = f"{prefix}.L{index:03d}"
        if page.get("chapter"):
            chapter = doc.add_paragraph()
            chapter.paragraph_format.space_before = Pt(10)
            chapter.paragraph_format.space_after = Pt(4)
            chapter_run = chapter.add_run(page["chapter"])
            set_run_font(chapter_run, size=11, color=ACCENT, bold=True)
        add_id_text_paragraph(
            doc,
            stable_id,
            page.get("text", ""),
            speaker=page.get("speaker"),
            kind=page.get("kind"),
        )
    doc.add_page_break()


def add_weekly_narratives(doc, weekly):
    doc.add_heading("週別共通ストーリー", level=1)
    intro = doc.add_paragraph(
        "週行動の前後に入る共通文です。キャラクター固有イベントとは別に、ミミ自身の物語を支える文章として確認してください。"
    )
    intro.paragraph_format.space_after = Pt(12)

    action_labels = {
        "work": "働く",
        "play": "遊ぶ",
        "rest": "休む",
        "search": "探す",
    }
    for week, entry in enumerate(weekly, start=1):
        doc.add_heading(f"第{week}週　{entry.get('title', '')}", level=2)
        add_metadata(doc, f"WEEK.{week:02d}.TITLE", "週タイトル", entry.get("title", ""))
        for index, text in enumerate(entry.get("setup", []), start=1):
            add_id_text_paragraph(
                doc,
                f"WEEK.{week:02d}.SETUP.L{index:02d}",
                text,
                kind="thought",
            )
        for action, text in entry.get("actions", {}).items():
            add_metadata(
                doc,
                f"WEEK.{week:02d}.ACTION.{action.upper()}",
                action_labels.get(action, action),
                text,
            )
    doc.add_page_break()


def add_fighters(doc, fighters):
    doc.add_heading("キャラクター固有ストーリー", level=1)
    stage_labels = {
        "meet": "出会い",
        "join": "加入",
        "bond": "すれ違い・関係深化",
        "power": "力の発現",
        "crisis": "危機・クライマックス",
        "liberation": "解放",
        "epilogue": "エピローグ",
    }
    for fighter_index, fighter in enumerate(fighters):
        if fighter_index > 0:
            doc.add_page_break()
        doc.add_heading(
            f"{fighter['name']}　｜　{fighter.get('kind', '')}",
            level=1,
        )
        prefix = f"CHAR.{fighter['id']}"
        add_metadata(doc, f"{prefix}.PROFILE.NAME", "名前", fighter["name"])
        add_metadata(doc, f"{prefix}.PROFILE.KIND", "種族・分類", fighter.get("kind", ""))
        add_metadata(doc, f"{prefix}.PROFILE.SUMMARY", "概要", fighter.get("summary", ""))
        add_metadata(
            doc,
            f"{prefix}.PROFILE.CURRENT_LIMIT",
            "現在の制約",
            fighter.get("currentLimit", ""),
        )
        add_metadata(
            doc,
            f"{prefix}.PROFILE.TRAIT_NAME",
            "特性名",
            fighter.get("traitName", ""),
        )
        add_metadata(
            doc,
            f"{prefix}.PROFILE.TRAIT_TEXT",
            "特性説明",
            fighter.get("traitText", ""),
        )

        for scene in fighter.get("scenes", []):
            stage = scene.get("stage", "scene")
            scene_title = scene.get("title", "")
            scene_copy = dict(scene)
            scene_copy["title"] = f"{stage_labels.get(stage, stage)}　｜　{scene_title}"
            add_scene(doc, scene_copy, f"{prefix}.{stage.upper()}", heading_level=2)


def add_matches_and_routes(doc, data):
    doc.add_page_break()
    doc.add_heading("大会・ルート文章", level=1)

    for group_name, group_key, group_prefix in (
        ("通常公式戦", "officialMatches", "MATCH.NORMAL"),
        ("支配ルート査定戦", "dominationMatches", "MATCH.DOMINATION"),
    ):
        doc.add_heading(group_name, level=2)
        for match in data[group_key]:
            prefix = f"{group_prefix}.{match['id']}"
            doc.add_heading(match.get("name", match["id"]), level=3)
            add_metadata(doc, f"{prefix}.NAME", "大会名", match.get("name", ""))
            add_metadata(
                doc,
                f"{prefix}.OPPONENT",
                "対戦相手名",
                match.get("opponentName", ""),
            )
            add_metadata(doc, f"{prefix}.STORY", "大会説明", match.get("story", ""))

    doc.add_heading("ルート紹介", level=2)
    for route in data.get("routes", []):
        route_id = route.get("id", "route")
        doc.add_heading(route.get("name", route_id), level=3)
        for key, value in route.items():
            if isinstance(value, str):
                add_metadata(
                    doc,
                    f"ROUTE.{route_id}.{key.upper()}",
                    key,
                    value,
                )


def add_source_note(doc, data):
    doc.add_page_break()
    doc.add_heading("編集後の受け渡し", level=1)
    paragraph = doc.add_paragraph()
    paragraph.add_run(
        "この文書を上書き編集して渡してください。灰色のIDが残っていれば、変更箇所を自動的に照合できます。"
        "IDを消してしまった箇所も前後関係から確認できますが、反映時の確認作業が増えます。"
    )
    add_metadata(doc, "EXPORT.META.SCHEMA", "抽出形式", data.get("schemaVersion", ""))
    add_metadata(doc, "EXPORT.META.DATE", "抽出日時", data.get("exportedAt", ""))


def main():
    data = json.loads(SOURCE.read_text(encoding="utf-8"))
    doc = Document()
    configure_document(doc)
    doc.core_properties.title = f"{data['projectTitle']} 全物語テキスト 編集用マスター"
    doc.core_properties.subject = "ゲーム内物語テキストの編集・レビュー"
    doc.core_properties.author = "Mimi Secret Boss Arena Project"

    add_cover(doc, data)
    add_prologue(doc, data["prologue"]["full"], "PROLOGUE.FULL", "プロローグ　通常版")
    add_prologue(
        doc,
        data["prologue"]["condensed"],
        "PROLOGUE.CONDENSED",
        "プロローグ　周回スキップ版",
    )

    doc.add_heading("導入・共通イベント", level=1)
    for scene in data["openingScenes"]:
        scene_key = scene["id"].replace(".", "_").upper()
        add_scene(doc, scene, f"OPENING.{scene_key}", heading_level=2)
    doc.add_page_break()

    add_weekly_narratives(doc, data["weeklyNarratives"])
    add_fighters(doc, data["fighters"])
    add_matches_and_routes(doc, data)
    add_source_note(doc, data)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
